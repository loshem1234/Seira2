"""seira_web.filegen — she produces files: md, docx, pdf, code.

Generated files are Corpus content (Art. 13: wholly temporal work
product, not identity) saved under corpus/outputs/, with a manifest
alongside the reference-file store's pattern. Each generation is
audit-logged so it shows up the same way any other authored act does.

Deliberately NOT supported here: image generation (needs a separate,
un-chosen third-party vendor — a real cost/provider decision, not a
default to pick silently) and rich docx styling beyond headings/
paragraphs/lists (a real but much larger scope; this covers the
common case honestly rather than faking full Word fidelity).
"""

from __future__ import annotations

import datetime as _dt
import json
import os
import re
import secrets
from pathlib import Path
from typing import Any, Dict, List

from seira_core.paths import seira_home

MAX_CONTENT_CHARS = 300_000  # generous; catches accidental runaway generation

SUPPORTED_FORMATS = {"md", "docx", "pdf", "code"}

# Loose but useful extension mapping for the 'code' format.
CODE_EXTENSIONS = {
    "python": "py", "javascript": "js", "typescript": "ts", "html": "html",
    "css": "css", "json": "json", "bash": "sh", "shell": "sh", "sql": "sql",
    "yaml": "yaml", "markdown": "md", "text": "txt",
}


class FileGenError(Exception):
    pass


def _outputs_dir() -> Path:
    return seira_home() / "corpus" / "outputs"


def _index_path() -> Path:
    return _outputs_dir() / "index.json"


def _load_index() -> Dict[str, Dict[str, Any]]:
    if not _index_path().exists():
        return {}
    return json.loads(_index_path().read_text(encoding="utf-8"))


def _save_index(index: Dict[str, Dict[str, Any]]) -> None:
    _outputs_dir().mkdir(parents=True, exist_ok=True)
    tmp = _index_path().with_suffix(".tmp")
    tmp.write_text(json.dumps(index, indent=2, ensure_ascii=False), encoding="utf-8")
    os.replace(tmp, _index_path())


def _now() -> str:
    return _dt.datetime.now(_dt.timezone.utc).isoformat()


def _safe_stem(filename: str) -> str:
    stem = re.sub(r"[^\w\-. ]", "", filename).strip() or "output"
    stem = re.sub(r"\.[A-Za-z0-9]{1,5}$", "", stem)  # drop any trailing extension
    return stem[:80] or "output"


def _write_markdown(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_code(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_docx(path: Path, content: str) -> None:
    """Minimal, honest Markdown-ish -> docx: '# ' headings, '- ' bullets,
    blank-line-separated paragraphs. Not full Markdown fidelity — real
    scope, stated plainly, not faked as complete."""
    from docx import Document
    doc = Document()
    for block in content.split("\n\n"):
        block = block.strip("\n")
        if not block.strip():
            continue
        lines = block.split("\n")
        first = lines[0]
        if first.startswith("### "):
            doc.add_heading(first[4:].strip(), level=3)
            lines = lines[1:]
        elif first.startswith("## "):
            doc.add_heading(first[3:].strip(), level=2)
            lines = lines[1:]
        elif first.startswith("# "):
            doc.add_heading(first[2:].strip(), level=1)
            lines = lines[1:]
        if not lines:
            continue
        if all(l.strip().startswith(("- ", "* ")) for l in lines if l.strip()):
            for l in lines:
                if l.strip():
                    doc.add_paragraph(l.strip()[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph("\n".join(lines).strip())
    doc.save(str(path))


def _write_pdf(path: Path, content: str) -> None:
    """Minimal, honest text/heading/bullet PDF via reportlab's Platypus.
    Same scope limitation as docx: structure, not full typesetting."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    story = []
    for block in content.split("\n\n"):
        block = block.strip("\n")
        if not block.strip():
            continue
        lines = block.split("\n")
        first = lines[0]
        if first.startswith("### "):
            story.append(Paragraph(escape(first[4:].strip()), styles["Heading3"]))
            lines = lines[1:]
        elif first.startswith("## "):
            story.append(Paragraph(escape(first[3:].strip()), styles["Heading2"]))
            lines = lines[1:]
        elif first.startswith("# "):
            story.append(Paragraph(escape(first[2:].strip()), styles["Heading1"]))
            lines = lines[1:]
        remaining = [l for l in lines if l.strip()]
        if remaining and all(l.strip().startswith(("- ", "* ")) for l in remaining):
            items = [ListItem(Paragraph(escape(l.strip()[2:].strip()), styles["Normal"]))
                     for l in remaining]
            story.append(ListFlowable(items, bulletType="bullet"))
        elif remaining:
            story.append(Paragraph(escape(" ".join(remaining)), styles["Normal"]))
        story.append(Spacer(1, 10))
    if not story:
        story = [Paragraph("(empty document)", styles["Normal"])]
    SimpleDocTemplate(str(path), pagesize=LETTER).build(story)


_WRITERS = {"md": _write_markdown, "docx": _write_docx, "pdf": _write_pdf, "code": _write_code}
_EXTENSIONS = {"md": "md", "docx": "docx", "pdf": "pdf"}  # 'code' resolved separately


def create_file(fmt: str, filename: str, content: str, language: str = "") -> Dict[str, Any]:
    if fmt not in SUPPORTED_FORMATS:
        raise FileGenError(f"format must be one of {sorted(SUPPORTED_FORMATS)}.")
    if not content.strip():
        raise FileGenError("Cannot create an empty file.")
    if len(content) > MAX_CONTENT_CHARS:
        raise FileGenError(f"Content too large ({len(content)} chars; "
                           f"limit {MAX_CONTENT_CHARS}).")
    out_id = f"out-{secrets.token_hex(6)}"
    ext = CODE_EXTENSIONS.get(language.lower(), "txt") if fmt == "code" else _EXTENSIONS[fmt]
    stem = _safe_stem(filename)
    disk_name = f"{out_id}.{ext}"
    _outputs_dir().mkdir(parents=True, exist_ok=True)
    _WRITERS[fmt](_outputs_dir() / disk_name, content)

    index = _load_index()
    record = {
        "out_id": out_id, "filename": f"{stem}.{ext}", "format": fmt,
        "language": language if fmt == "code" else None,
        "disk_name": disk_name, "created": _now(),
        "size": (_outputs_dir() / disk_name).stat().st_size,
    }
    index[out_id] = record
    _save_index(index)
    return record


def list_outputs() -> List[Dict[str, Any]]:
    return sorted(_load_index().values(), key=lambda r: r["created"], reverse=True)


def get_output_path(out_id: str) -> Path:
    index = _load_index()
    rec = index.get(out_id)
    if rec is None:
        raise FileGenError(f"No generated file {out_id!r}.")
    return _outputs_dir() / rec["disk_name"]


def get_output_record(out_id: str) -> Dict[str, Any]:
    index = _load_index()
    rec = index.get(out_id)
    if rec is None:
        raise FileGenError(f"No generated file {out_id!r}.")
    return rec
